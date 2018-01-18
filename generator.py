#!/usr/bin/python
# -*- coding: utf-8 -*-
from docx import Document

from tkinter import *
from tkinter import ttk

import tkSimpleDialog as tkSD
import tkFileDialog
import tkMessageBox

import legal

import os


class mainapp:

    def __init__(self, master, *args, **kwargs):
        self.master = master
        master.title("Affidavit Generator")

        # top menubar
        self.menubar = Menu(root)
        self.filemenu = Menu(self.menubar, tearoff=0)
        self.filemenu.add_command(label="About", command=self.about)
        self.filemenu.add_separator()
        self.filemenu.add_command(label="Exit", command=root.quit)
        self.menubar.add_cascade(label="File", menu=self.filemenu)
        root.config(menu=self.menubar)

        # frame for showing what info is needed
        self.title = Label(root, text="What you'll need:",
                           bg=mycolor, fg='white')
        self.title.grid(row=2, column=0, padx=4, pady=4, sticky="NW")

        self.info_frame = Frame(root, highlightbackground=bordercolor,
                                highlightcolor=bordercolor,
                                highlightthickness=1, bd=0)
        self.info_frame.grid(row=2, column=1, padx=4, pady=4, sticky="W")

        self.info = StringVar()
        self.info_label = Label(self.info_frame, wraplength=200,
                                bg=buttoncolor, fg='white',
                                textvariable=self.info, anchor="w")
        self.info_label.grid(row=2, column=1, sticky="W")

        # combobox for selecting affidavit
        self.label = Label(root, text="Select the type of affidavit:",
                           padx=4, pady=4, bg=mycolor, fg='white')
        self.label.grid(row=1, column=0, sticky="NW")

        self.value = StringVar()
        self.afftype = ttk.Combobox(self.master, textvariable=self.value,
                                    state='readonly')
        self.afftype['values'] = aff_info_required.keys()
        self.afftype.grid(row=1, column=1, sticky="W")
        self.afftype.bind('<<ComboboxSelected>>', self.displayinfo)
        self.afftype.current(0)
        self.displayinfo(self)

        # button for running python-docx script
        self.make = Button(root, text="Generate", fg='white', bg=buttoncolor,
                           command=self.make_doc)
        self.make.grid(row=7, column=0, padx=10, pady=10)

        # directory select button
        self.save_location = Button(root, text='Browse', fg='white',
                                    bg=buttoncolor, command=self.set_directory)
        self.save_location.grid(row=5, column=3, padx=10, pady=10)

        self.save_label = Label(root, text="Save location:",
                                bg=mycolor, fg='white')
        self.save_label.grid(row=5, column=0, padx=4, pady=4, sticky="W")

        self.save_directory = StringVar()
        self.save_directory.set(os.getcwd())
        self.info_label = Label(root, wraplength=200,
                                bg=buttoncolor, fg='white',
                                textvariable=self.save_directory,
                                anchor="w").grid(row=5, column=1, sticky="W")

    def client_exit(self):
        exit()

    def set_directory(self):
        tmp_dir = tkFileDialog.askdirectory()
        self.save_directory.set(tmp_dir)

    def get_directory(self):
        current_dir = self.save_directory.get()
        return current_dir

    def about(self):
            tkMessageBox.showinfo("About",
                                  "made by Shaun Anderson with python-docx")

    def displayinfo(self, event):
        print "User selection is %s" % self.afftype.get()
        x = []
        for i in aff_info_required[self.afftype.get()]:
            x.append(i)
        self.info.set('\n'.join(x))

    def get_info(self):

            doc_type = self.afftype.get()

            self.studentName = tkSD.askstring("Student Name",
                                              "Please enter the student's name")

            self.spouseName = tkSD.askstring("Spouse Name",
                                             "Please enter the spouse's name")
            self.location = tkSD.askstring("Location",
                                           "Please enter the city and province")

            if doc_type == 'Common Law':
                self.livingsince = tkSD.askstring("Date Living Since",
                                                 ("Please enter the date the "
                                                  "couple began living together"))
                self.separated = ''

            elif doc_type == 'Separated':
                self.livingsince = ''
                self.separated = tkSD.askstring("Date Separated",
                                               ("Please enter "
                                                "the date the couple separated"))

            self.children = tkMessageBox.askquestion("Children",
                                                     "Do they have dependents?")

            self.childrenNames = []
            self.childrenBdays = []

            if self.children == 'yes':
                self.number_of_children = tkSD.askinteger("Number",
                                                     "How many kids?")
                for i in range(0, self.number_of_children):
                    self.childrenNames.append(tkSD.askstring("Name", legal.q_chn
                                                        % (i + 1)))

                    self.childrenBdays.append(tkSD.askstring("BDAY", (legal.q_ch
                                                        % (i + 1))))
            else:
                self.number_of_children = 0

            if doc_type == 'Separated':

                self.open_cst_window()
                self.st_custody = self.app.selection

                # allows user to select custody details
                # if self.shrd == 'yes':
                #
                #    self.open_cst_window()
                #    self.st_custody = self.app.cst_results
                #    self.open_spcst_window()
                #    self.sp_custody = self.app.cst_results

            return {'Student Name':
                    self.studentName,
                    'Spouse Name':
                    self.spouseName,
                    'Location':
                    self.location,
                    'Living Since':
                    self.livingsince,
                    'Children Names':
                    self.childrenNames,
                    'Children Birthdays':
                    self.childrenBdays,
                    'Children':
                    self.number_of_children,
                    'Custody':
                    self.st_custody,
                    'Separated Date':
                    self.separated
                    }

    def open_cst_window(self):

        self.cst_window = Toplevel(self.master)
        self.app = Custody(self.cst_window)
        self.master.wait_window(self.cst_window)

    def make_doc(self):
        # grab user selected directory and affidavit type
        self.get_directory()
        self.answers = self.get_info()
        print self.answers
        document = Document('affidavit-template.docx')
        document.save('affidavit-%s.docx' % self.answers["Student Name"])
        document._body.clear_content()

        # affidavit header with scilicet
        title = document.add_heading('AFFIDAVIT')
        title.style = document.styles['Heading 1']
        run = title.add_run()
        run.add_break()

        country = document.add_table(rows=4, cols=2)
        country.style = 'none'

        canada = country.cell(0, 0)
        canada.text = 'Canada'

        rightBracket = country.cell(0, 1)
        rightBracket.text = ')'

        province = country.cell(1, 0)
        province.text = 'Province of Ontario'

        provinceRB = country.cell(1, 1)
        provinceRB.text = ') S. S.'

        finalRB = country.cell(2, 1)
        finalRB.text = ')'

        if self.afftype.get() == 'Separated':
            p = document.add_paragraph(legal.sole_declare %
                                       (self.answers["Student Name"],
                                        self.answers["Location"]))
            run = p.add_run()
            run.add_break()

        elif self.afftype.get() == 'Common Law':
            p = document.add_paragraph(legal.declare %
                                       (self.answers["Student Name"],
                                        self.answers["Spouse Name"],
                                        self.answers["Location"]))
            run = p.add_run()
            run.add_break()

        # declaration
        p = document.add_paragraph(legal.aff_purpose)
        p.style = document.styles['ListNumber']
        run = p.add_run()
        run.add_break()

        if self.afftype.get() == 'Common Law':

            p = document.add_paragraph(legal.living % self.answers["Living Since"])
            p.style = document.styles['ListNumber']
            run = p.add_run()
            run.add_break()

        elif self.afftype.get() == 'Separated':
            p = document.add_paragraph(legal.separated)
            p.style = document.styles['ListNumber']
            run = p.add_run()
            run.add_break()

            p = document.add_paragraph("The name of my spouse is %s" self.answers["Spouse Name"])
            p.style = document.styles['ListNumber']
            run = p.add_run()
            run.add_break()

            p = document.add_paragraph(legal.living % self.answers["Separated Date"])
            p.style = document.styles['ListNumber']
            run = p.add_run()
            run.add_break()

        if self.answers["Children"] > 0:

            p = document.add_paragraph(legal.custodial % self.answers['Children'])
            p.style = document.styles['ListNumber']
            run = p.add_run()
            run.add_break()
            run.add_break()

            children_names = self.answers['Children Names']
            children_bdays = self.answers['Children Birthdays']

            for i in range(0, self.answers["Children"]):
                    run = p.add_run("%s, born %s" %
                                    (children_names[i], children_bdays[i]))
                    run.add_break()
                    run.add_break()

        if self.afftype.get() == 'Separated':
            if self.answers['Custody'] == 'Sole Support':

                p = document.add_paragraph(legal.solecust)
                p.style = document.styles['ListNumber']
                run = p.add_run()
                run.add_break()


        p = document.add_paragraph(legal.marriage_law)
        p.style = document.styles['ListNumber']

        run = p.add_run()
        run.add_break()
        run.add_break()
        run.add_break()

        # table to hold signatures
        table = document.add_table(rows=6, cols=2)
        table.style = 'Hello'

        sworn = table.cell(0, 0)
        sworn.text = legal.affirmed

        firstUnderline = table.cell(0, 1)
        firstUnderline.text = legal.underline

        firstSig = table.cell(0, 1)
        firstSig.add_paragraph(self.answers["Student Name"])

        commUnderline = table.cell(5, 0)
        commUnderline.add_paragraph(legal.underline)

        commissioner = table.cell(5, 0)
        commissioner.add_paragraph("Mark Alexander Robinson")

        secondUnderline = table.cell(5, 1)
        secondUnderline.add_paragraph(legal.underline)

        secondSig = table.cell(5, 1)
        secondSig.add_paragraph(self.answers["Spouse Name"])

        filename = str("Affidavit-%s.docx" % self.answers["Student Name"])
        filedirectory = self.get_directory()
        filepath = os.path.join(filedirectory, filename)

        document.save(filepath)


class Custody:
    def __init__(self, master):
        self.master = master
        master.title("Custody Details")
        self.msg = Message(master,
                           text="Is the custody shared or does the student have sole support")
        self.msg.grid(row=0, column=0)

        # create listbox of children names entered
        self.student_custody = Listbox(master, selectmode=EXTENDED)
        self.student_custody.grid(row=1, column=0)

        self.student_custody.insert(END, "Shared Custody")
        self.student_custody.insert(END, "Sole Support")

        self.btn = Button(master, text="Okay", command=self.select)
        self.btn.grid(row=2, column=0)

    def select(self):

        self.selection = self.student_custody.get(self.student_custody.curselection())
        self.master.destroy()


if __name__ == '__main__':

    # ui colors
    mycolor = '#2c2f33'
    buttoncolor = '#4f555c'
    bordercolor = '#99aab5'
    root = Tk()
    root.configure(bg=mycolor)
    aff_info_required = {'Common Law':
                         ['Student Name', 'Spouse Name',
                          'Number of Children', 'Names of Children',
                          'Birthdates of Children', 'Location',
                          'Date living since'],
                         'Separated':
                         ['Student Name', 'Spouse Name',
                          'Number of Children', 'Names of Children',
                          'Birthdates of Children', 'Location',
                          'Custody Details', 'Address']
                         }

    app = mainapp(root)
    root.geometry('450x300')
    root.mainloop()
